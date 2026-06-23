import io
import logging
from typing import Optional
from fastapi import HTTPException
from backend.services.translation_engine import translate_text

log = logging.getLogger("FreeTranslate")

# ── Optional document translation imports ──────────────────────────────────
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False


def translate_docx(file_bytes: bytes, source_lang: str, target_lang: str) -> bytes:
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx not installed.")
    doc = DocxDocument(io.BytesIO(file_bytes))
    
    # Process paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            translated, _, _ = translate_text(para.text, source_lang, target_lang)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = translated
            else:
                para.add_run(translated)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        translated, _, _ = translate_text(para.text, source_lang, target_lang)
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = translated
                        else:
                            para.add_run(translated)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def translate_pdf(file_bytes: bytes, source_lang: str, target_lang: str) -> bytes:
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=500, detail="pdfplumber not installed.")
    pages_text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            pages_text.append(page.extract_text() or "")

    translated_pages = []
    for page_text in pages_text:
        if page_text.strip():
            t, _, _ = translate_text(page_text, source_lang, target_lang)
            translated_pages.append(t)
        else:
            translated_pages.append("")

    # Reconstruct as plain-text PDF
    if FPDF_AVAILABLE:
        pdf_out = FPDF()
        pdf_out.set_auto_page_break(auto=True, margin=15)
        pdf_out.add_page()
        pdf_out.set_font("Helvetica", size=11)
        for page_text in translated_pages:
            for line in page_text.split("\n"):
                pdf_out.multi_cell(0, 8, line)
            pdf_out.add_page()
        buf = io.BytesIO()
        pdf_out.output(buf)
        return buf.getvalue()
    else:
        # Fallback: plain text
        return "\n\n---PAGE BREAK---\n\n".join(translated_pages).encode("utf-8")
